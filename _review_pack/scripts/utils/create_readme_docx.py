#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""README.docx（Word版）を作成するスクリプト"""

import sys
import io
from pathlib import Path

# Windows環境でのコンソール出力の文字化けを防ぐ
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_readme_docx(output_path: Path):
    """README.docxを作成"""
    doc = Document()
    
    # ページ設定
    section = doc.sections[0]
    section.page_width = Cm(21)  # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    
    # タイトル
    title = doc.add_heading('ippi-down（イッピダウン）', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('入札情報ファイル自動ダウンローダー')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.bold = True
    
    doc.add_paragraph()
    
    manual_title = doc.add_paragraph('操作マニュアル')
    manual_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    manual_title_run = manual_title.runs[0]
    manual_title_run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # 概要
    intro = doc.add_paragraph()
    intro.add_run('このソフトウェアは、建設情報サービス「i-ppi.jp」から入札情報ファイルを自動的にダウンロードするためのツールです。')
    
    doc.add_paragraph()
    
    # 重要なお知らせ（枠囲み風）
    notice = doc.add_paragraph()
    notice.add_run('★ 重要 ★').bold = True
    doc.add_paragraph('・Pythonなどのプログラミング環境は不要です。')
    doc.add_paragraph('・このフォルダ内の ippi-down.exe をダブルクリックするだけで使用できます。')
    
    doc.add_page_break()
    
    # 目次
    doc.add_heading('目次', level=1)
    toc_items = [
        '1. はじめに（動作環境）',
        '2. 起動方法',
        '3. 基本的な使い方',
        '4. 画面の説明',
        '5. よくある質問（FAQ）',
        '6. トラブルシューティング',
        '7. フォルダ構成',
        '8. お問い合わせ',
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # 1. はじめに
    doc.add_heading('1. はじめに（動作環境）', level=1)
    
    doc.add_heading('必要な環境', level=2)
    doc.add_paragraph('・Windows 10 または Windows 11（64ビット版）')
    doc.add_paragraph('・インターネット接続')
    
    doc.add_heading('不要なもの', level=2)
    doc.add_paragraph('・Python環境（インストール不要）')
    doc.add_paragraph('・管理者権限（不要）')
    doc.add_paragraph('・特別なソフトウェア（不要）')
    
    doc.add_paragraph()
    
    # 2. 起動方法
    doc.add_heading('2. 起動方法', level=1)
    
    doc.add_heading('手順', level=2)
    steps = [
        'このフォルダ内にある「ippi-down.exe」を探します',
        '「ippi-down.exe」をダブルクリックします',
        'アプリケーションが起動します',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('初回起動時の注意', level=2)
    doc.add_paragraph('・初回起動時は、起動に5〜10秒程度かかることがあります')
    doc.add_paragraph('・Windows Defenderが警告を表示する場合があります')
    
    notice_para = doc.add_paragraph()
    notice_para.add_run('　→ 「詳細情報」をクリックし、「実行」を選択してください').italic = True
    
    notice_para2 = doc.add_paragraph()
    notice_para2.add_run('　→ これは正常な動作です（PyInstallerで作成されたexeファイルの特性です）').italic = True
    
    doc.add_paragraph()
    
    # 3. 基本的な使い方
    doc.add_heading('3. 基本的な使い方', level=1)
    
    doc.add_heading('ステップ1：検索条件を設定する', level=2)
    doc.add_paragraph('画面左側の「検索条件」エリアで条件を入力します')
    doc.add_paragraph('・発注機関：ダウンロードしたい発注機関を選択')
    doc.add_paragraph('・工事名：キーワードを入力（例：「トンネル」「橋梁」）')
    doc.add_paragraph('・日付範囲：必要に応じて期間を指定')
    
    doc.add_heading('ステップ2：保存先を指定する', level=2)
    doc.add_paragraph('・画面の「保存先」欄で、ファイルを保存するフォルダを指定します')
    doc.add_paragraph('・「参照」ボタンをクリックしてフォルダを選択できます')
    doc.add_paragraph('・指定しない場合は、このフォルダ内の「downloads」フォルダに保存されます')
    
    doc.add_heading('ステップ3：ダウンロードを開始する', level=2)
    doc.add_paragraph('・「ダウンロード開始」ボタンをクリックします')
    doc.add_paragraph('・進捗バーでダウンロード状況を確認できます')
    doc.add_paragraph('・完了すると通知が表示されます')
    
    doc.add_heading('ダウンロードを中止したい場合', level=2)
    doc.add_paragraph('・「キャンセル」ボタンをクリックすると、ダウンロードを中止できます')
    doc.add_paragraph('・既にダウンロード済みのファイルはそのまま保存されます')
    
    doc.add_page_break()
    
    # 4. 画面の説明
    doc.add_heading('4. 画面の説明', level=1)
    
    doc.add_paragraph('アプリケーションの画面は以下の要素で構成されています：')
    doc.add_paragraph()
    
    # テーブルで画面要素を説明
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    headers = ['画面要素', '説明']
    cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cells[i].text = header
        cells[i].paragraphs[0].runs[0].bold = True
    
    data = [
        ('検索条件エリア', '発注機関、工事名、日付範囲などの検索条件を入力します'),
        ('保存先', 'ダウンロードしたファイルの保存先フォルダを指定します'),
        ('ダウンロード開始ボタン', 'クリックするとダウンロードが開始されます'),
        ('キャンセルボタン', 'ダウンロード中にクリックすると処理を中止します'),
        ('進捗バー', 'ダウンロードの進捗状況を表示します'),
    ]
    
    for i, (element, desc) in enumerate(data, 1):
        cells = table.rows[i].cells
        cells[0].text = element
        cells[1].text = desc
    
    doc.add_paragraph()
    
    # 5. FAQ
    doc.add_heading('5. よくある質問（FAQ）', level=1)
    
    faqs = [
        ('Q1. Pythonをインストールする必要がありますか？', 
         'A1. いいえ、不要です。このexeファイルだけで動作します。'),
        ('Q2. 管理者権限は必要ですか？', 
         'A2. いいえ、不要です。通常のユーザー権限で動作します。'),
        ('Q3. ダウンロードしたファイルはどこに保存されますか？', 
         'A3. 画面で指定した保存先フォルダに保存されます。指定しない場合は、このフォルダ内の「downloads」フォルダに保存されます。'),
        ('Q4. 設定は保存されますか？', 
         'A4. はい、保存されます。次回起動時に前回の設定が読み込まれます。設定は「config/config.yaml」ファイルに保存されています。'),
        ('Q5. 定期的に自動実行できますか？', 
         'A5. はい、スケジュール機能があります。「設定」ボタンから、実行間隔と時刻を設定できます。'),
        ('Q6. 同じファイルを何度もダウンロードしてしまいますか？', 
         'A6. いいえ、既にダウンロード済みのファイルは自動的にスキップされます。'),
    ]
    
    for q, a in faqs:
        q_para = doc.add_paragraph()
        q_para.add_run(q).bold = True
        doc.add_paragraph(a)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # 6. トラブルシューティング
    doc.add_heading('6. トラブルシューティング', level=1)
    
    troubles = [
        ('アプリケーションが起動しない', 
         'Windows Defenderがブロックしている可能性があります',
         ['Windows Defenderの警告が出た場合は「詳細情報」→「実行」を選択',
          'それでも起動しない場合は、ウイルス対策ソフトの設定を確認してください']),
        ('「〇〇が見つかりません」というエラーが出る', 
         '必要なファイルが不足している可能性があります',
         ['フォルダ構成が正しいか確認してください',
          '再度ダウンロードしてお試しください']),
        ('ダウンロードが途中で止まる', 
         'ネットワーク接続の問題、またはサーバーの一時的な問題',
         ['インターネット接続を確認してください',
          'しばらく待ってから再度お試しください',
          '自動リトライ機能があるため、一時的なエラーは自動で回復します']),
    ]
    
    for problem, cause, solutions in troubles:
        doc.add_heading(f'【問題】{problem}', level=2)
        doc.add_paragraph(f'原因：{cause}')
        doc.add_paragraph('対処：')
        for i, sol in enumerate(solutions, 1):
            doc.add_paragraph(f'　{i}. {sol}')
        doc.add_paragraph()
    
    doc.add_heading('ログファイルの確認方法', level=2)
    doc.add_paragraph('詳細なエラー情報は「logs/app.log」ファイルに記録されています。')
    doc.add_paragraph('メモ帳などのテキストエディタで開いて確認できます。')
    
    doc.add_page_break()
    
    # 7. フォルダ構成
    doc.add_heading('7. フォルダ構成', level=1)
    
    folder_table = doc.add_table(rows=8, cols=2)
    folder_table.style = 'Table Grid'
    
    folder_data = [
        ('ファイル/フォルダ', '説明'),
        ('ippi-down.exe', '実行ファイル（これをダブルクリック）'),
        ('README.txt', '操作マニュアル（テキスト版）'),
        ('README.docx', '操作マニュアル（Word版・このファイル）'),
        ('config/', '設定ファイルフォルダ'),
        ('downloads/', 'ダウンロードしたファイルの保存先（自動生成）'),
        ('logs/', 'ログファイルフォルダ（自動生成）'),
    ]
    
    for i, (name, desc) in enumerate(folder_data):
        cells = folder_table.rows[i].cells
        cells[0].text = name
        cells[1].text = desc
        if i == 0:
            cells[0].paragraphs[0].runs[0].bold = True
            cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run('※ config.yaml、downloads/、logs/ は初回起動時に自動的に作成されます').italic = True
    
    doc.add_paragraph()
    
    # 8. お問い合わせ
    doc.add_heading('8. お問い合わせ', level=1)
    doc.add_paragraph('本ソフトウェアに関するお問い合わせは、管理者までご連絡ください。')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # バージョン情報
    doc.add_heading('バージョン情報', level=1)
    
    version_table = doc.add_table(rows=4, cols=2)
    version_table.style = 'Table Grid'
    
    version_data = [
        ('ソフトウェア名', 'ippi-down（入札情報ファイルダウンローダー）'),
        ('バージョン', '2.1'),
        ('対応OS', 'Windows 10/11（64ビット）'),
        ('作成日', '2026年1月'),
    ]
    
    for i, (label, value) in enumerate(version_data):
        cells = version_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
    
    # 保存
    doc.save(output_path)
    print(f'README.docxを作成しました: {output_path}')


def main():
    root = Path(__file__).parent.parent.parent  # scripts/utils/ から プロジェクトルートへ
    output_path = root / 'release' / 'ippi-down-dist' / 'README.docx'
    
    print('README.docx（Word版）を作成中...')
    create_readme_docx(output_path)
    print('完了！')


if __name__ == '__main__':
    main()
